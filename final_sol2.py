import os
import re
import json
from docx import Document

#list of lines that are prefixes/suffixes that are not needed and only intrupt
INVALID_LINES = [
    'מנהל הוועדה:', 'רשמת פרלמנטרית:', 'מוזמנים באמצעים מקוונים:', 'חברי כנסת:', 'יועצת משפטית:',
    'מנהל/ת הוועדה:', 'חברי הוועדה:', 'חברי הכנסת:', 'משתתפים באמצעים מקוונים:', 'מנהלת הוועדה:',
    'סדר היום:', 'משתתפים (באמצעים מקוונים):', 'רישום פרלמנטרי:',
    ' החדשים בנוגע לחפירות בהר הבית – הצעה לסדר של חברי הכנסת:', 'יועץ משפטי:', 'קצרנית פרלמנטרית:',
    'קצרנית:', 'רצח הירקן בטירה ופעולות משטרת ישראלהיו"ר ש\' וייס:', '(23 בינואר 2008):', '(7 ביוני 2006)',
    'אי-אמון בממשלה בשל:', 'הצעות חוק:', 'רשמה וערכה:', 'סדר-היום:', 'נוכחים:', 'ייעוץ משפטי:', 'מוזמנים:',
    'נכחו:'
]


HEB_LETTERS = list("אבגדהוזחטיכלמנסעפצקרשתץףךםן")
ENG_LETTERS = list("abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ")

#common patterns like date, acronyms, splitting patterns
DATE_PATTERNS = [
    r"\d{1,2}/\d{1,2}(?:/\d{2,4})",
    r"\d{1,2}\.\d{1,2}(?:\.\d{2,4})"
]
TIME_PATTERN = r"\b(?:[0-1]?\d|2[0-3]):[0-5]\d\b"
ACRONYM_PATTERN = r'\w+"\w+'
NUMBER_PATTERN = r'(\d)+(,\d{3})*(\.\d+)?(%|$)?'
TOKEN_SPLIT_PATTERN = r"([^א-ת0-9'])"

HEBREW_ABC = list("אבגדהוזחטיכלמנסעפצקרשתץףךםן")

def hebrew_words_to_number(text):
    units = {
        "אחד": 1, "אחת": 1, "הראשונה": 1,
        "שניים": 2, "שתיים": 2, "שני": 2, "שתים": 2, "השנייה": 2,
        "שלושה": 3, "שלוש": 3, "השלישית": 3,
        "ארבעה": 4, "ארבע": 4, "הרביעית": 4,
        "חמישה": 5, "חמש": 5, "החמישית": 5,
        "שישה": 6, "שש": 6, "השישית": 6,
        "שבעה": 7, "שבע": 7, "השביעית": 7,
        "שמונה": 8, "השמינית": 8,
        "תשעה": 9, "תשע": 9, "התשיעית": 9
    }
    special = {
        "האחד-עשר": 11, "האחת-עשר": 11,
        "השנים-עשר": 12, "השתים-עשרה": 12,
        "השלושה-עשר": 13, "השלוש-עשרה": 13,
        "הארבעה-עשר": 14, "הארבע-עשרה": 14,
        "החמישה-עשר": 15, "החמש-עשרה": 15,
        "השישה-עשר": 16, "השש-עשרה": 16,
        "השבעה-עשר": 17, "השבע-עשרה": 17,
        "השמונה-עשר": 18, "השמונה-עשרה": 18,
        "התשעה-עשר": 19, "התשעה-עשרה": 19
    }
    tens = {
        "עשרה": 10, "עשר": 10, "העשר": 10,
        "עשרים": 20, "העשרים": 20,
        "שלושים": 30, "השלושים": 30,
        "ארבעים": 40, "הארבעים": 40,
        "חמישים": 50, "החמישים": 50,
        "שישים": 60, "השישים": 60,
        "שבעים": 70, "השבעים": 70,
        "שמונים": 80, "השמונים": 80,
        "תשעים": 90, "התשעים": 90
    }
    hundreds = {
        "מאה": 100, "המאה": 100,
        "מאתיים": 200, "המאתיים": 200,
        "שלוש מאות": 300, "השלוש מאות": 300,
        "ארבע מאות": 400, "הארבע מאות": 400,
        "חמש מאות": 500, "החמש מאות": 500,
        "שש מאות": 600, "השש מאות": 600,
        "שבע מאות": 700, "השבע מאות": 700,
        "שמונה מאות": 800, "השמונה מאות": 800,
        "תשע מאות": 900, "התשע מאות": 900
    }

    if text in special:
        return special[text]

    words = text.replace("-", " ").split()
    total = 0
    i = 0
    #parse phrase to sum it up
    while i < len(words):
        w = words[i]
        #remove the "and" prefix
        if w.startswith("ו") and len(w) > 1:
            w = w[1:]
        if i + 1 < len(words):
            combined = f"{words[i]} {words[i+1]}"
            if combined in hundreds:
                total += hundreds[combined]
                i += 2
                continue
        if w in units:
            total += units[w]
        elif w in tens:
            total += tens[w]
        elif w in hundreds:
            total += hundreds[w]
        i += 1
    return total if total > 0 else -1

def extract_protocol_number(doc, protocol_type):
    if protocol_type == "committee":
        for para in doc.paragraphs:
            text = para.text.strip()
            match_digits = re.search(r'(?:פרוטוקול|הפרוטוקול|ישיבה|הישיבה)(?:\s+(?:מס[\'’]?|מספר))?\s*([0-9]+)', text)
            if match_digits:
                return int(match_digits.group(1))
    else:
        # plenary number often appears in hebrew after "הישיבה"
        for para in doc.paragraphs:
            text = para.text.strip()
            if "הישיבה" in text:
                match = re.search(r"הישיבה\s+([\u0590-\u05FF\-]+)", text)
                if match:
                    num = hebrew_words_to_number(match.group(1))
                    if num > 0:
                        return num
    return -1

def is_line_excludable(text):
    #check if line is in our list of known irrelevant lines
    return any(invalid in text for invalid in INVALID_LINES)

#check the style of paragraph (bold or underlined)
def check_underline(para):

    if para.style and ((para.style.base_style and para.style.base_style.font and para.style.base_style.font.underline) or
                       (para.style.font and para.style.font.underline)):
        return True
    return any(r.underline for r in para.runs)

def check_bold(para):
    if para.style and ((para.style.base_style and para.style.base_style.font and para.style.base_style.font.bold) or
                       (para.style.font and para.style.font.bold)):
        return True
    return any(r.bold for r in para.runs)

def extract_name(name):
    # remove all unneeded things in the speaker name (paranetheses, prefixes, suffixes)
    if name.endswith(':'):
        name = name[:-1].strip()

    titles = [
        '<<.*>>', 'תשובת', 'הד"ר', 'ד"ר', 'מ"מ היו"ר', 'היו"ר', 'יו"ר', 'נשיא הפרלמנט האירופי', 'שר', 'שרת',
        'עו"ד', 'המשנה לראש הממשלה', 'משנה לראש הממשלה', 'ראש הממשלה', 'נצ"מ', 'מר', 'ניצב', 'טפסר משנה', 'רשף',
        'פרופ\'', 'סגן', 'סגנית', 'השר', 'השרה', 'מזכיר', 'מזכירת', 'ועדת'
    ]
    departments = [
        'הכנסת', 'הכלכלה', 'הכלכלה והתכנון', 'האנרגיה והמים', 'החינוך', 'החינוך והתרבות',
        'התשתיות הלאומיות', 'התשתיות הלאומיות, האנרגיה והמים', 'חינוך', 'המודיעין', 'לשיתוף פעולה אזורי',
        'ההסברה והתפוצות', 'האוצר', 'להגנת הסביבה', 'התעשייה, המסחר והתעסוקה', 'לאיכות הסביבה',
        'התרבות והספורט', 'התשתיות', 'המשפטים', 'הרווחה', 'הרווחה והשירותים החברתיים', 'הבינוי והשיכון',
        'התחבורה והבטיחות בדרכים', 'המשטרה', 'הבריאות', 'החקלאות ופיתוח הכפר', 'התקשורת',
        'במשרד ראש הממשלה', 'הפנים', 'הביטחון', 'לביטחון פנים', 'המדע והטכנולוגיה', 'העלייה והקליטה',
        'לקליטת העלייה', 'התיירות', 'החוץ', 'המדע,', 'למודיעין', 'לאזרחים ותיקים', 'לענייני דתות', 'לענייני מודיעין',
        'המדע', 'התרבות', 'הספורט', 'תרבות', 'ספורט'
    ]


    patterns = titles + departments
    prefix_pattern = r'^(' + '|'.join(map(re.escape, patterns)) + r')\s+'
    while True:
        new_name = re.sub(prefix_pattern, '', name).strip()
        if new_name == name:
            break
        name = new_name

    name_parts = name.split()
    for i, token in enumerate(name_parts):
        if token.strip() in patterns:
            name_parts[i] = ""
    cleaned = " ".join(t for t in name_parts if t.strip())
    #handle paranetheses
    cleaned = re.sub(r'\(.*?\)', '', cleaned).strip()
    return cleaned

def is_speaker_line(text, paragraph, protocol_type):
    #check if the paragraph is speaker
    if not text or is_line_excludable(text):
        return False
    if protocol_type == "committee":
        return (not check_bold(paragraph)) and check_underline(paragraph) and text.endswith(':')
    else:
        return check_bold(paragraph) and check_underline(paragraph) and text.endswith(':')

def check_valid_sentence(sentence):
    # check the sentence: must have hebrew chars, no english chars , and no " --- " patterns
    if re.search(r"((-|–)\s*){2}", sentence):
        return False
    contains_hebrew = any(char in HEB_LETTERS for char in sentence)
    contains_english = any(char in ENG_LETTERS for char in sentence)
    return contains_hebrew and not contains_english

def split_sentence_to_tokens(sentence_text):
    # split sentence to tokens by using regex
    tokens = []
    def split_word(word):
        if not word:
            return []
        # check for known patterns in text and split them if found
        for patterns in [DATE_PATTERNS, [TIME_PATTERN], [NUMBER_PATTERN], [ACRONYM_PATTERN]]:
            for pattern in patterns:
                m = re.search(pattern, word)
                if m:
                    return split_word(word[:m.start()]) + [m.group()] + split_word(word[m.end():])
        # if not found then split by TOKEN_SPLIT_PATTERN for punctuation
        return [p for p in re.split(TOKEN_SPLIT_PATTERN, word) if p]

    for w in sentence_text.split():
        tokens.extend(split_word(w))
    tokens = [t.strip() for t in tokens if t.strip()]
    return tokens

def sentence_splitter(text):
    #this splits the sentence
    endings = ['.', '!', '?', '\n', ';']
    sentences = []
    start_index = 0
    for i, c in enumerate(text):
        if c in endings:
            if start_index != i:
                # dont split on floating point numbers like 69.42
                if c == '.' and i + 1 < len(text) and text[i+1].isdigit():
                    continue
                # next char should be space/Hebrew/newline
                if i + 1 < len(text) and text[i+1] not in HEB_LETTERS + [' ', '\n']:
                    continue
                candidate = text[start_index:i + 1].strip()
                if check_valid_sentence(candidate):
                    tokenized = split_sentence_to_tokens(candidate)
                    if len(tokenized) >= 4:
                        sentences.append(" ".join(tokenized))
            start_index = i + 1
    return sentences

def extract_type_from_filename(file_name):
    # find kneset num and protocol type from filename
    knesset_num = int(file_name[:2])
    proto_type = "committee" if file_name[5] == 'v' else "plenary"
    return knesset_num, proto_type

def process_document(doc_path, protocol_type):
    # parse docx file
    doc = Document(doc_path)
    number_protocol = extract_protocol_number(doc, protocol_type)
    results = []
    current_speaker = None
    current_text = ""
    for p in doc.paragraphs:
        p_text = p.text.strip()
        if is_speaker_line(p_text, p, protocol_type):
            # we read a speaker and need to collect the text related to this speaker
            if current_speaker and current_text:
                speaker_clean = extract_name(current_speaker)
                for s in sentence_splitter(current_text):
                    results.append((speaker_clean, s))
            current_speaker = p_text
            current_text = ""
        else:
            if current_speaker and p_text:
                current_text += p_text + " "
    # last speaker is special case
    if current_speaker and current_text:
        speaker_clean = extract_name(current_speaker)
        for s in sentence_splitter(current_text):
            results.append((speaker_clean, s))
    return number_protocol, results

if __name__ == "__main__":
    folder_with_files_path = "TextFiles"
    output_file_name = "result.jsonl"

    if not os.path.isdir(folder_with_files_path):
        print("specified folder path does not exist")
    else:
        docx_files = [f for f in os.listdir(folder_with_files_path) if f.endswith('.docx')]
        if not docx_files:
            print("no docx files found")
        else:
            # use set to avoid duplicate rows
            all_sentences = []
            seen = set()
            for file_name in docx_files:
                print(f"now working on: {file_name}")
                knesset_number, protocol_type = extract_type_from_filename(file_name)
                doc_path = os.path.join(folder_with_files_path, file_name)
                number_protocol, speaker_sentences = process_document(doc_path, protocol_type)

                #only save the unique rows
                for (speaker_nm, text_sentence) in speaker_sentences:
                    row_key = (file_name, knesset_number, protocol_type, number_protocol if number_protocol is not None else -1, speaker_nm, text_sentence)
                    if row_key not in seen:
                        seen.add(row_key)
                        entry = {
                            "name_protocol": file_name,
                            "number_knesset": knesset_number,
                            "type_protocol": protocol_type,
                            "number_protocol": number_protocol if number_protocol is not None else -1,
                            "name_speaker": speaker_nm,
                            "text_sentence": text_sentence
                        }
                        all_sentences.append(entry)
            #write to json lines
            with open(output_file_name, "w", encoding="utf-8") as out_f:
                for e in all_sentences:
                    out_f.write(json.dumps(e, ensure_ascii=False) + "\n")

            print("completed")
