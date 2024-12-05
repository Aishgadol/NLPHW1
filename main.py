import os
import re as re
from docx import Document

hebrew_abc = ['א', 'ב', 'ג', 'ד', 'ה', 'ו', 'ז', 'ח', 'ט', 'י', 'כ', 'ל', 'מ', 'נ', 'ס', 'ע', 'פ', 'צ', 'ק', 'ר', 'ש',
             'ת', 'ץ', 'ף', 'ך', 'ם', 'ן']

english_abc = ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v',
       'w', 'x', 'y', 'z', 'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R',
       'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']

bad_lines = ['מנהל הוועדה:', 'רשמת פרלמנטרית:', 'מוזמנים באמצעים מקוונים:', 'חברי כנסת:', 'יועצת משפטית:',
                 'מנהל/ת הוועדה:', 'חברי הוועדה:', 'חברי הכנסת:', 'משתתפים באמצעים מקוונים:', 'מנהלת הוועדה:',
                 'סדר היום:', 'משתתפים (באמצעים מקוונים):', 'רישום פרלמנטרי:',
                 ' החדשים בנוגע לחפירות בהר הבית – הצעה לסדר של חברי הכנסת:', 'יועץ משפטי:', 'קצרנית פרלמנטרית:',
                 'קצרנית:', 'רצח הירקן בטירה ופעולות משטרת ישראלהיו"ר ש\' וייס:', '(23 בינואר 2008):', '(7 ביוני 2006)',
                 'אי-אמון בממשלה בשל:', 'הצעות חוק:', 'רשמה וערכה:', 'סדר-היום:', 'נוכחים:', 'ייעוץ משפטי:', 'מוזמנים:',
                 'נכחו:']

lines=[]
filenames_data={}
prot_type=""
directory = "TextFiles"
filenames=[f for f in os.listdir(directory) if f.endswith('.docx')]
my_corp=[]
curr_speaker=""
curr_speaker_text=""
speaker_flag=False

tot_sent={}


def extract_filename_data():
    for i in range(len(filenames)):
        #knesset num could be 1 or 2 digits

        tokens=filenames[i].split("_")
        tokens[2]=tokens[2][:tokens[2].index('.')]
        if tokens[1] == "ptv":
            prot_type = "committee"
        elif tokens[1] == "ptm":
            prot_type = "plenary"
        else:
            raise ValueError("Unexpected Error, Dont know ",tokens[1])
        filenames_data[i]=(int(tokens[0]),prot_type,int(tokens[2]))

def hebrew_words_to_number(text):
    units = {
        "אחד": 1, "אחת": 1, "הראשונה": 1,
        "שניים": 2, "שתיים": 2, "שני": 2, "שתים": 2, "השנייה": 2,
        "שלושה": 3, "שלוש": 3, "השלישית": 3,
        "ארבעה": 4, "ארבע": 4, "הרביעית": 4,
        "חמישה": 5, "חמש": 5, "החמישית": 5,
        "שישה": 6, "שש": 6, "השישית": 6,
        "שבעה": 7, "שבע": 7, "השביעית": 7,
        "שמונה": 8, "השמינית": 8,
        "תשעה": 9, "תשע": 9, "התשיעית": 9
    }

    special = {
        "האחד-עשר": 11, "האחת-עשר": 11,
        "השנים-עשר": 12, "השתים-עשרה": 12,
        "השלושה-עשר": 13, "השלוש-עשרה": 13,
        "הארבעה-עשר": 14, "הארבע-עשרה": 14,
        "החמישה-עשר": 15, "החמש-עשרה": 15,
        "השישה-עשר": 16, "השש-עשרה": 16,
        "השבעה-עשר": 17, "השבע-עשרה": 17,
        "השמונה-עשר": 18, "השמונה-עשרה": 18,
        "התשעה-עשר": 19, "התשעה-עשרה": 19
    }

    tens = {
        "עשרה": 10, "עשר": 10, "העשר": 10,
        "עשרים": 20, "העשרים": 20,
        "שלושים": 30, "השלושים": 30,
        "ארבעים": 40, "הארבעים": 40,
        "חמישים": 50, "החמישים": 50,
        "שישים": 60, "השישים": 60,
        "שבעים": 70, "השבעים": 70,
        "שמונים": 80, "השמונים": 80,
        "תשעים": 90, "התשעים": 90
    }

    hundreds = {
        "מאה": 100, "המאה": 100,
        "מאתיים": 200, "המאתיים": 200,
        "שלוש מאות": 300, "השלוש מאות": 300,
        "ארבע מאות": 400, "הארבע מאות": 400,
        "חמש מאות": 500, "החמש מאות": 500,
        "שש מאות": 600, "השש מאות": 600,
        "שבע מאות": 700, "השבע מאות": 700,
        "שמונה מאות": 800, "השמונה מאות": 800,
        "תשע מאות": 900, "התשע מאות": 900
    }

    if text in special:
        return special[text]

    # Clean and split the text
    words = text.replace("-", " ").split()
    total = 0
    i = 0

    while i < len(words):
        word = words[i]

        if word.startswith("ו") and len(word) > 1:
            word = word[1:]

        if i + 1 < len(words):
            combined = f"{words[i]} {words[i + 1]}"
            if combined in hundreds:
                total += hundreds[combined]
                i += 2
                continue

        if word in units:
            total += units[word]
        elif word in tens:
            total += tens[word]
        elif word in hundreds:
            total += hundreds[word]
        elif word == "ו":
            pass

        i += 1

    return total if total > 0 else -1

def check_underline(para):
    # check if the sentence is underlined
    # there is 3 ways - 1. in style.base_style.font.underline 2. style.font.underline
    if para.style is not None and ((para.style.base_style is not None and para.style.base_style.font.underline) or (
            para.style.font is not None and para.style.font.underline)):
        return True
    # 3. run.underline : we go through all the runs because the first one isn't necessarily enough
    elif para.runs:
        for run in para.runs:
            if run.underline:
                return True
    return False

def check_bold(para):
    # check if the sentence is boldd
    # there is 3 ways - 1. in style.base_style.font.bold 2. style.font.bold
    if para.style is not None and ((para.style.base_style is not None and para.style.base_style.font.bold) or (
            para.style.font is not None and para.style.font.bold)):
        return True
    # 3. run.bold : we go through all the runs because the first one isn't necessarily enough
    elif para.runs:
        for run in para.runs:
            if run.bold:
                return True
    return False

def extract_name(name):

    titles = [
        '<<.*>>',
        'תשובת',
        'הד"ר', 'ד"ר', 'מ"מ היו"ר', 'היו"ר', 'יו"ר',
        'נשיא הפרלמנט האירופי','שר','שרת',
        'עו"ד', 'המשנה לראש הממשלה', 'משנה לראש הממשלה', 'ראש הממשלה',
        'נצ"מ', 'מר', 'ניצב', 'טפסר משנה', 'רשף', 'פרופ\'',
        'סגן', 'סגנית', 'השר', 'השרה', 'מזכיר', 'מזכירת', 'ועדת',
    ]
    departments = [
    'הכנסת', 'הכלכלה', 'הכלכלה והתכנון', 'האנרגיה והמים', 'החינוך', 'החינוך והתרבות',
    'התשתיות הלאומיות', 'התשתיות הלאומיות, האנרגיה והמים','חינוך',
    'המודיעין', 'לשיתוף פעולה אזורי', 'ההסברה והתפוצות', 'האוצר', 'להגנת הסביבה',
    'התעשייה, המסחר והתעסוקה', 'לאיכות הסביבה', 'התרבות והספורט', 'התשתיות',
    'המשפטים', 'הרווחה', 'הרווחה והשירותים החברתיים', 'הבינוי והשיכון',
    'התחבורה והבטיחות בדרכים', 'המשטרה', 'הבריאות', 'החקלאות ופיתוח הכפר',
    'התקשורת', 'במשרד ראש הממשלה', 'הפנים', 'הביטחון', 'לביטחון פנים',
    'המדע והטכנולוגיה', 'העלייה והקליטה', 'לקליטת העלייה', 'התיירות', 'החוץ','המדע,',
    'למודיעין', 'לאזרחים ותיקים', 'לענייני דתות', 'לענייני מודיעין','המדע','התרבות','הספורט','תרבות','ספורט'
    ]
    patterns=titles + departments
    pattern = r'^(' + '|'.join(map(re.escape, patterns)) + r')\s+'
    while True:
        new_name = re.sub(pattern, '', name).strip()
        if new_name == name:
            break
        name = new_name
    tokens=new_name.split(" ")
    for i in range(len(tokens)):
        if tokens[i].strip() in patterns:
            tokens[i]=""
    new_name=" ".join(str(s) for s in tokens if s)
    return new_name.strip()

def extract_protocol_number(index,file_path):
    file = Document(file_path)
    try:
        if filenames_data[index][1] == "committee":

            for para in file.paragraphs:
                text = para.text.strip()

                match_digits = re.search(r'(?:פרוטוקול|הפרוטוקול|ישיבה|הישיבה)(?:\s+(?:מס[\'’]?|מספר))?\s*([0-9]+)',
                                         text)
                if match_digits:
                    return match_digits.group(1)  # Extract numeric value


        else:

            for paragraph in file.paragraphs:
                text = paragraph.text.strip()
                if "הישיבה" in text:
                    match = re.search(r"הישיבה\s+([\u0590-\u05FF\-]+)", text)
                    if match:
                        hebrew_number_phrase = match.group(1)
                        return str(hebrew_words_to_number(hebrew_number_phrase))


    except Exception as e:
        print(f"Error processing x file {filenames[index]}: {e}")

def turn_text_to_sentences(text):
    output=[]
    start_idx=0
    endings = ['.' , '!' , '?' , '\n' , ';' , '\t']
    #go through each character in the text
    for index,char in enumerate(text):
        #are we looking at an ending character?
        if char in endings:
            if start_idx != index:
                #this case handles floating point numbers
                if char=='.' and index+1<len(text) and text[index+1].isdigit():
                    continue
                #this case handles quotes
                if (index+1<len(text)) and (text[index+1] not in hebrew_abc) and (text[index+1]!=' ') and text[index+1]!='\n':
                    continue


                #now we can extract the sentence
                sentence=text[start_idx:index+1].strip()
                if (sentence and sentence !="\n"):
                    output.append(sentence)
            start_idx = index+1
    return output



extract_filename_data()

# for m in filenames_data:
#     print(m," : ",filenames_data[m])


# count=0
# for i in range(len(filenames)):
#     doc=Document(os.path.join(directory,filenames[i]))
#     for para in doc.paragraphs:
#         if ("!!" in para.text) or ("??" in para.text) or ("!?" in para.text) or ("?!" in para.text) or ("..." in para.text) or (".." in para.text):
#             print("BINGO IN "+filenames[i]+"\n"+para.text+"\n"+"-"*30)
#             count+=1
#
# print(f"total {count} bingos")
# exit(0)
sent_count=0
for i in range(len(filenames)):
    curr_speaker_text=""
    curr_speaker=""
    doc=Document(os.path.join(directory,filenames[i]))
    doc_names_set=[]
    protocol_name=filenames[i]
    knesset_number=filenames_data[i][0]
    protocol_type=filenames_data[i][1]
    #filenames_data[i][2] is fileNumber, useless. annoying.
    protocol_number=extract_protocol_number(i,os.path.join(directory,filenames[i]))


    #print(protocol_name,"  ,  ",knesset_number,"  ,  ",protocol_type,"  ,  ",protocol_number)
    dealing_with_comittee=protocol_type=="committee"
    for paragraph in doc.paragraphs:
        if(dealing_with_comittee):
            #if the next condition is true, this means we're looking at a speaker
            if((not check_bold(paragraph)) and check_underline(paragraph) and paragraph.text.endswith(':')):

                #check the speaker is not a single word, would be error
                if(len(paragraph.text.split(" "))<2 or paragraph.text in bad_lines):
                    continue

                #if next condition is true, means that we've already encounterd a speaker and we need to save all the text
                #that the speaker said, and we need to seperate them into sentences
                if(curr_speaker):
                   speaker_sentences=turn_text_to_sentences(curr_speaker_text)
                   print(f'------------------------\n '
                         f'for speaker : {curr_speaker} in file: {filenames[i]}, sentences in paragraph: ')
                   for sentence in speaker_sentences:
                       #my_corp.append((protocol_name,knesset_number,prot_type,protocol_number,curr_speaker,sentence))
                       print(sentence)
                       tot_sent[sentence]=filenames[i]
                       sent_count+=1
                   curr_speaker_text=""
                   print("-"*50+"\n")
                #no else is needed cuz if we're looking at a speaker, this means that the former speaker
                #has done talking.
                curr_speaker=""
                curr_speaker = extract_name(paragraph.text.strip()[:-1])

            #if this condition is true, this means we're still on the same speaker and we need to
            #accumulate the text his taking.
            elif (not curr_speaker==""):
                    curr_speaker_text += paragraph.text
            else:
                continue


        elif(not dealing_with_comittee): #not dealing with comittee so its plenary maybe

            #if the next condition is true, this means we're looking at a speaker
            if(check_bold(paragraph) and check_underline(paragraph) and paragraph.text.endswith(':')):

                #check the speaker is not a single word, would be error
                if(len(paragraph.text.split(" "))<2 or paragraph.text in bad_lines):
                    continue

                #if next condition is true, means that we've already encounterd a speaker and we need to save all the text
                #that the speaker said, and we need to seperate them into sentences
                if(curr_speaker):
                    speaker_sentences=turn_text_to_sentences(curr_speaker_text)
                    print(f'------------------------\n '
                          f'for speaker : {curr_speaker} in file: {filenames[i]}, sentences in paragraph: ')
                    for sentence in speaker_sentences:
                        #my_corp.append((protocol_name,knesset_number,prot_type,protocol_number,curr_speaker,sentence))
                        print(sentence)
                        tot_sent[sentence]=filenames[i]
                        sent_count+=1
                    curr_speaker_text=""
                    print("-"*50+"\n")
                #no else is needed cuz if we're looking at a speaker, this means that the former speaker
                #has done talking.
                curr_speaker=""
                curr_speaker = extract_name(paragraph.text.strip()[:-1])

            #if this condition is true, this means we're still on the same speaker and we need to
            #accumulate the text his taking.
            elif (not curr_speaker==""):
                curr_speaker_text += paragraph.text
            else:
                continue

        else:
            print("never shouldve come here")

    curr_speaker=""
    curr_speaker_text=""
    print ("\n"+"-"*30+"\n")

print(f'total {sent_count} sentences counted\n number of sentences of size 3,2,1:\n{sum(1 for s in tot_sent if len(s.split(" ")) < 4)}')

print(f'five shortest sentences: \n')
for i,x in enumerate(sorted(tot_sent,key=len)[:100]):
    print(i,x)
    print(tot_sent[x],"\n")
print("--"*40)
print(f'five longest sentences: \n')
for x in sorted(tot_sent,key=len)[-5:]:
    print(x)
    print(tot_sent[x],"\n")
print("--"*40)